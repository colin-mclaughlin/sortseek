import os
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional
from fastapi import UploadFile
import pdfplumber
from docx import Document as DocxDocument
import logging
from sqlalchemy.orm import Session
import hashlib
from datetime import datetime

from database import get_db
from models import Document
from services.search_service import SearchService

logger = logging.getLogger(__name__)

DOCUMENTS_DIR = Path("data/documents")

class FileService:
    """Service for handling file operations and document parsing"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    def __init__(self, search_service: SearchService = None):
        self.data_dir = DOCUMENTS_DIR
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.search_service = search_service
        if self.search_service is None:
            logger.warning("FileService initialized without SearchService - indexing will be disabled")
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file contents"""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating hash for {file_path}: {e}")
            return None
    
    def _get_file_modified_time(self, file_path: Path) -> float:
        """Get file's last modified timestamp"""
        try:
            return file_path.stat().st_mtime
        except Exception as e:
            logger.error(f"Error getting modified time for {file_path}: {e}")
            return None
    
    def _check_file_changed(self, file_path: Path, existing_doc: Document) -> bool:
        """Check if file has changed since last import/indexing"""
        try:
            # Get current file stats
            current_size = file_path.stat().st_size
            current_modified_time = self._get_file_modified_time(file_path)
            current_hash = self._calculate_file_hash(file_path)
            
            # Compare with stored values
            size_changed = existing_doc.file_size != current_size
            time_changed = existing_doc.modified_time != current_modified_time
            hash_changed = existing_doc.file_hash != current_hash
            
            # Log the comparison for debugging
            logger.info(f"File change check for {file_path.name}:")
            logger.info(f"  Size: {existing_doc.file_size} -> {current_size} (changed: {size_changed})")
            logger.info(f"  Modified time: {existing_doc.modified_time} -> {current_modified_time} (changed: {time_changed})")
            logger.info(f"  Hash: {existing_doc.file_hash} -> {current_hash} (changed: {hash_changed})")
            
            # File is considered changed if any of these differ
            return size_changed or time_changed or hash_changed
            
        except Exception as e:
            logger.error(f"Error checking file changes for {file_path}: {e}")
            # If we can't check, assume file has changed to be safe
            return True

    async def import_folder(self, folder_path: str, force: bool = False) -> Dict[str, Any]:
        """Import all supported documents from a folder recursively"""
        folder = Path(folder_path)
        if not folder.exists():
            raise ValueError(f"Folder {folder_path} does not exist")
        
        imported_files = []
        skipped_files = []
        indexed_files = []
        
        # Count total files for progress tracking
        total_files = sum(1 for file_path in folder.rglob("*") 
                         if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS)
        
        logger.info(f"Starting folder import: {folder_path} ({total_files} files, force={force})")
        
        # Walk through all files in the folder
        for file_path in folder.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc_dict = await self._import_single_file(file_path, force=force)
                    if doc_dict:
                        imported_files.append(doc_dict)
                        
                        # Check if file was actually indexed or skipped
                        db = next(get_db())
                        doc = db.query(Document).filter(Document.id == doc_dict['id']).first()
                        if doc and doc.is_indexed:
                            indexed_files.append(doc_dict['filename'])
                        else:
                            skipped_files.append(doc_dict['filename'])
                            
                except Exception as e:
                    logger.error(f"Failed to import {file_path}: {e}")
                    continue
        
        # Log summary
        logger.info(f"Folder import completed: {folder_path}")
        logger.info(f"  Total files processed: {len(imported_files)}")
        logger.info(f"  Files indexed: {len(indexed_files)}")
        logger.info(f"  Files skipped (unchanged): {len(skipped_files)}")
        
        if skipped_files:
            logger.info(f"  Skipped files: {', '.join(skipped_files[:10])}{'...' if len(skipped_files) > 10 else ''}")
        
        return {
            "imported_files": imported_files,
            "indexed_count": len(indexed_files),
            "skipped_count": len(skipped_files),
            "total_count": len(imported_files),
            "indexed_files": indexed_files,
            "skipped_files": skipped_files
        }
    
    async def import_file(self, file: UploadFile) -> Dict[str, Any]:
        """Import a single uploaded file"""
        temp_path = DOCUMENTS_DIR / file.filename
        with open(temp_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        try:
            doc_dict = await self._import_single_file(temp_path)
            return doc_dict
        finally:
            if temp_path.exists():
                temp_path.unlink()
    
    async def _extract_content(self, file_path: Path) -> Optional[str]:
        """Extract text content from different file types"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return await self._extract_pdf_content(file_path)
            elif extension == '.docx':
                return await self._extract_docx_content(file_path)
            elif extension == '.txt':
                return await self._extract_txt_content(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return None
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from PDF file and return joined text (for legacy use)"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                content = '\n'.join(text_parts)
                print(f"📄 Extracted content from PDF: {file_path.name} → {len(content)} chars")
                return content
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return ""

    async def extract_pdf_chunks(self, file_path: Path) -> list:
        """Extract text content from PDF file, split by page, and return list of chunks"""
        try:
            with pdfplumber.open(file_path) as pdf:
                chunks = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        chunks.append((i+1, text.strip()))
                return chunks
        except Exception as e:
            logger.error(f"Error extracting PDF chunks: {e}")
            return []

    async def _import_single_file(self, file_path: Path, force: bool = False) -> Optional[Dict[str, Any]]:
        """Import a single file and extract its content, then index by chunk if PDF"""
        try:
            # Check if file already exists in database
            db = next(get_db())
            existing_doc = db.query(Document).filter(Document.file_path == str(file_path)).first()
            
            if existing_doc:
                logger.info(f"File {file_path} already exists in database")
                
                # Check if file has changed (unless force=True)
                if not force:
                    file_changed = self._check_file_changed(file_path, existing_doc)
                    if not file_changed:
                        logger.info(f"File {file_path.name} unchanged - skipping reindexing")
                        return existing_doc.to_dict()
                    else:
                        logger.info(f"File {file_path.name} has changed - reindexing")
                else:
                    logger.info(f"Force reindexing requested for {file_path.name}")
                
                # File has changed or force=True, so we need to reindex
                # Extract new content
                content = await self._extract_content(file_path)
                if not content:
                    logger.warning(f"Could not extract content from {file_path}")
                    return existing_doc.to_dict()  # Return existing doc even if extraction fails
                
                # Update file metadata
                current_size = file_path.stat().st_size
                current_modified_time = self._get_file_modified_time(file_path)
                current_hash = self._calculate_file_hash(file_path)
                
                # Update the existing document
                existing_doc.content = content
                existing_doc.file_size = current_size
                existing_doc.modified_time = current_modified_time
                existing_doc.file_hash = current_hash
                existing_doc.is_indexed = False  # Mark for reindexing
                existing_doc.embedding_path = None  # Clear old embedding path
                existing_doc.updated_at = datetime.utcnow()
                
                db.commit()
                db.refresh(existing_doc)
                
                logger.info(f"Updated existing document: {file_path}")
                document = existing_doc
                
            else:
                # New file - extract content and create new document
                logger.info(f"New file detected: {file_path}")
                content = await self._extract_content(file_path)
                if not content:
                    logger.warning(f"Could not extract content from {file_path}")
                    return None
                
                # Calculate file metadata
                current_size = file_path.stat().st_size
                current_modified_time = self._get_file_modified_time(file_path)
                current_hash = self._calculate_file_hash(file_path)
                
                # Create new document record
                document = Document(
                    filename=file_path.name,
                    file_path=str(file_path),
                    file_type=file_path.suffix.lower(),
                    file_size=current_size,
                    content=content,
                    is_indexed=False,
                    imported_at=datetime.utcnow(),
                    modified_time=current_modified_time,
                    file_hash=current_hash
                )
                db.add(document)
                db.commit()
                db.refresh(document)
                logger.info(f"Successfully imported new file: {file_path}")
            
            # Index the document based on file type
            if self.search_service:
                file_type = file_path.suffix.lower()
                if file_type == '.pdf':
                    # For PDFs, index by page chunks
                    page_chunks = await self.extract_pdf_chunks(file_path)
                    if page_chunks:
                        chunk_texts = [text for (page, text) in page_chunks]
                        chunk_metadatas = [{
                            "document_id": document.id,
                            "filename": document.filename,
                            "file_path": document.file_path,
                            "file_type": document.file_type,
                            "page": page
                        } for (page, text) in page_chunks]
                        await self.search_service.index_document_chunks(document, chunk_texts, chunk_metadatas)
                elif file_type in ['.docx', '.txt']:
                    # For DOCX and TXT, index the entire document content
                    await self.search_service.index_document(document)
            
            return document.to_dict()
            
        except Exception as e:
            logger.error(f"Error importing {file_path}: {e}")
            return None
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            content = '\n'.join(text_parts)
            print(f"📄 Extracted content from DOCX: {file_path.name} → {len(content)} chars")
            return content
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            return ""
    
    async def _extract_txt_content(self, file_path: Path) -> str:
        """Extract text content from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                print(f"📄 Extracted content from TXT: {file_path.name} → {len(content)} chars")
                return content
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                    print(f"📄 Extracted content from TXT (latin-1): {file_path.name} → {len(content)} chars")
                    return content
            except Exception as e:
                logger.error(f"Error reading TXT file with latin-1 encoding: {e}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting TXT content: {e}")
            return ""
    
    def get_document_by_id(self, document_id: int) -> Optional[Document]:
        """Get a document by its ID"""
        try:
            db = next(get_db())
            return db.query(Document).filter(Document.id == document_id).first()
        except Exception as e:
            logger.error(f"Error getting document {document_id}: {e}")
            return None
    
    def get_all_documents(self, skip: int = 0, limit: int = 100) -> List[Document]:
        """Get all documents with pagination"""
        try:
            db = next(get_db())
            return db.query(Document).offset(skip).limit(limit).all()
        except Exception as e:
            logger.error(f"Error getting documents: {e}")
            return [] 